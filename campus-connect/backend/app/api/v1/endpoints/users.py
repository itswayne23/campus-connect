from fastapi import APIRouter, HTTPException, status, Depends, Query, UploadFile, File
from typing import List, Optional
from app.schemas.user import UserUpdate, UserResponse, UserProfileResponse
from app.services.user_service import user_service
from app.services.follow_service import follow_service
from app.api.deps import get_current_user_id, get_optional_user_id

import uuid
import os
from io import BytesIO
from docx import Document

router = APIRouter()

@router.get("/search", response_model=List[UserResponse])
async def search_users(
    q: str = Query("", min_length=0),
    current_user_id: str = Depends(get_current_user_id)
):
    if not q or len(q.strip()) == 0:
        return []
    users = await user_service.search_users(q, current_user_id)
    return users

@router.get("/{user_id}", response_model=UserProfileResponse)
async def get_user_profile(
    user_id: str,
    current_user_id: Optional[str] = Depends(get_optional_user_id)
):
    user = await user_service.get_user_by_id(user_id, current_user_id)
    
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )
    
    return user

@router.put("/{user_id}", response_model=UserResponse)
async def update_user(
    user_id: str,
    user_data: UserUpdate,
    current_user_id: str = Depends(get_current_user_id)
):
    if user_id != current_user_id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not authorized to update this profile"
        )
    
    user = await user_service.update_user(user_id, user_data)
    
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )
    
    return user

@router.delete("/{user_id}")
async def delete_user(
    user_id: str,
    current_user_id: str = Depends(get_current_user_id)
):
    if user_id != current_user_id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Not authorized to delete this account"
        )
    
    success = await user_service.delete_user(user_id)
    
    if not success:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to delete user"
        )
    
    return {"message": "User deleted successfully"}

@router.get("/{user_id}/followers", response_model=List[UserResponse])
async def get_followers(
    user_id: str,
    current_user_id: str = Depends(get_current_user_id)
):
    followers = follow_service.get_followers(user_id, current_user_id)
    return followers

@router.get("/{user_id}/following", response_model=List[UserResponse])
async def get_following(
    user_id: str,
    current_user_id: str = Depends(get_current_user_id)
):
    following = follow_service.get_following(user_id, current_user_id)
    return following

@router.post("/{user_id}/follow")
async def follow_user(
    user_id: str,
    current_user_id: str = Depends(get_current_user_id)
):
    result = await follow_service.follow_user(current_user_id, user_id)
    return result

@router.delete("/{user_id}/follow")
async def unfollow_user(
    user_id: str,
    current_user_id: str = Depends(get_current_user_id)
):
    result = await follow_service.unfollow_user(current_user_id, user_id)
    return result

@router.post("/{user_id}/block")
async def block_user(
    user_id: str,
    current_user_id: str = Depends(get_current_user_id)
):
    if user_id == current_user_id:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Cannot block yourself"
        )
    
    from app.core import get_service_client
    service = get_service_client()
    
    try:
        service.table("blocks").insert({
            "blocker_id": current_user_id,
            "blocked_id": user_id
        }).execute()
        
        if follow_service._is_following(current_user_id, user_id):
            follow_service.unfollow_user(current_user_id, user_id)
        
        return {"message": "User blocked successfully"}
    except Exception as e:
        if "23505" in str(e):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="User already blocked"
            )
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to block user"
        )

@router.get("/{user_id}/posts")
async def get_user_posts(
    user_id: str,
    current_user_id: str = Depends(get_current_user_id)
):
    from app.services.post_service import post_service
    
    posts = await post_service.get_user_posts(user_id, current_user_id)
    return posts

@router.post("/{user_id}/report")
async def report_user(
    user_id: str,
    reason: str = Query(..., min_length=1),
    current_user_id: str = Depends(get_current_user_id)
):
    if user_id == current_user_id:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Cannot report yourself"
        )
    
    from app.core import get_service_client
    service = get_service_client()
    
    try:
        service.table("reports").insert({
            "reporter_id": current_user_id,
            "reported_id": user_id,
            "reason": reason
        }).execute()
        
        return {"message": "User reported successfully"}
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to report user"
        )

@router.post("/avatar/upload")
async def upload_avatar(
    file: UploadFile = File(...),
    current_user_id: str = Depends(get_current_user_id)
):
    from app.core import get_service_client
    
    if not file:
        raise HTTPException(status_code=400, detail="No file provided")
    
    file_ext = os.path.splitext(file.filename)[1].lower()
    allowed_images = ['.jpg', '.jpeg', '.png', '.gif', '.webp']
    
    if file_ext not in allowed_images:
        raise HTTPException(status_code=400, detail="File type not allowed. Allowed: jpg, jpeg, png, gif, webp")
    
    bucket_name = "avatars"
    file_name = f"{current_user_id}/{uuid.uuid4()}{file_ext}"
    
    file_content = await file.read()
    
    if len(file_content) > 5 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File size must be less than 5MB")
    
    service = get_service_client()
    
    try:
        response = service.storage.from_(bucket_name).upload(
            file_name,
            file_content,
            {"content-type": file.content_type}
        )
        
        public_url = service.storage.from_(bucket_name).get_public_url(file_name)
        
        service.table("profiles").update({"avatar_url": public_url}).eq("id", current_user_id).execute()
        
        return {"url": public_url}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@router.delete("/avatar")
async def delete_avatar(
    current_user_id: str = Depends(get_current_user_id)
):
    from app.core import get_service_client
    
    service = get_service_client()
    
    try:
        service.table("profiles").update({"avatar_url": None}).eq("id", current_user_id).execute()
        return {"message": "Avatar deleted successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to delete avatar: {str(e)}")

@router.get("/export-profile")
async def export_profile_docx(
    current_user_id: str = Depends(get_current_user_id)
):
    from app.core import get_service_client
    from fastapi.responses import StreamingResponse
    
    service = get_service_client()
    
    profile_response = service.table("profiles").select("*").eq("id", current_user_id).execute()
    
    if not profile_response.data:
        raise HTTPException(status_code=404, detail="Profile not found")
    
    profile = profile_response.data[0]
    
    doc = Document()
    doc.add_heading('User Profile', 0)
    
    doc.add_heading('Personal Information', level=1)
    doc.add_paragraph(f"Username: {profile.get('username', 'N/A')}")
    doc.add_paragraph(f"Email: {profile.get('email', 'N/A')}")
    doc.add_paragraph(f"University: {profile.get('university', 'N/A')}")
    doc.add_paragraph(f"Bio: {profile.get('bio', 'N/A')}")
    
    doc.add_heading('Statistics', level=1)
    doc.add_paragraph(f"Followers: {profile.get('followers_count', 0)}")
    doc.add_paragraph(f"Following: {profile.get('following_count', 0)}")
    doc.add_paragraph(f"Posts: {profile.get('posts_count', 0)}")
    
    doc.add_heading('Account Information', level=1)
    doc.add_paragraph(f"Member since: {profile.get('created_at', 'N/A')}")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={profile.get('username', 'profile')}_profile.docx"}
    )
